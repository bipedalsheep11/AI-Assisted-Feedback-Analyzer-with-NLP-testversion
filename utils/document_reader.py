# utils/file_reader.py

import fitz                    # pymupdf — for PDF files
from docx import Document      # python-docx — for Word files
from pathlib import Path
import pandas as pd
import ipywidgets as widgets
import re
import io
import pdfplumber
# ─────────────────────────────────────────────
#  PDF EXTRACTION
# ─────────────────────────────────────────────
# def extract_text_from_pdf(file) -> str:
#     """
#     Extract all text from a PDF file.

#     Parameters
#     ----------
#     file : str | Path | BytesIO
#         Either a file path (string or Path object) or a file-like object
#         (e.g., from Streamlit's st.file_uploader which returns BytesIO).

#     Returns
#     -------
#     str
#         The full extracted text of the PDF, with pages separated by
#         a newline and a page marker.
#     """
#     # fitz.open() accepts both file paths and bytes/BytesIO objects
#     file.seek(0)
#     doc = fitz.open(stream=file.read(), filetype="pdf")

#     pages_text = []
#     for page_number, page in enumerate(doc, start=1):
#         # extract_text() returns the text on one page as a single string
#         # The "text" flag preserves reading order (left-to-right, top-to-bottom)
#         page_text = page.get_text("text")
#         if page_text.strip():  # Skip blank pages
#             pages_text.append(f"[Page {page_number}]\n{page_text}")

#     doc.close()
#     return "\n\n".join(pages_text)

def extract_pages_from_pdf(file) -> list[dict]:
    """
    Extract text page-by-page. Returns a list of dicts so you can
    reference which page a piece of content came from — useful for citations.

    Returns
    -------
    list of {"page": int, "text": str}
    """
    file.seek(0)
    doc = fitz.open(stream=file.read(), filetype="pdf")
    pages = []
    for i, page in enumerate(doc, start=1):
        text = page.get_text("text").strip()
        if text:
            pages.append({"page": i, "text": text})
    doc.close()
    return pages

# #-------Extracting Content from Docx-------------
# def extract_text_from_docx(file) -> str:
#     """
#     Extract all text from a Word .docx file.

#     Parameters
#     ----------
#     file : str | Path | BytesIO
#         File path or file-like object from Streamlit uploader.

#     Returns
#     -------
#     str
#         Full document text. Headings are preserved with a marker
#         so the LLM can understand document structure.
#     """
#     doc = Document(file)

#     parts = []
#     for paragraph in doc.paragraphs:
#         text = paragraph.text.strip()
#         if not text:
#             continue  # Skip empty paragraphs

#         # Preserve heading structure — tells the LLM where sections begin
#         # paragraph.style.name is e.g. "Heading 1", "Heading 2", "Normal"
#         style = paragraph.style.name
#         if style.startswith("Heading"):
#             parts.append(f"\n## {text}")
#         else:
#             parts.append(text)

#     return "\n".join(parts)

# def extract_tables_from_docx(file) -> list[pd.DataFrame]:
#     """
#     Extract all tables from a Word document as pandas DataFrames.
#     CSF evaluation forms often use tables — this preserves their structure.

#     Returns
#     -------
#     list of pd.DataFrame, one per table found in the document.
#     """
#     doc = Document(file)
#     tables = []

#     for table in doc.tables:
#         rows = []
#         for row in table.rows:
#             # Each cell's text, stripped of whitespace
#             row_data = [cell.text.strip() for cell in row.cells]
#             rows.append(row_data)

#         if rows:
#             # First row treated as column headers if it looks like a header
#             df = pd.DataFrame(rows[1:], columns=rows[0])
#             tables.append(df)

#     return tables
    
# def extract_structured_responses_from_docx(file) -> pd.DataFrame:
#     """
#     Specialized extractor for evaluation forms where each row in a table
#     represents one respondent's answers.

#     Combines table extraction with text normalization to produce
#     a DataFrame ready for the sentiment/theme pipeline.
#     """
#     tables = extract_tables_from_docx(file)

#     if not tables:
#         # No tables found — treat the whole doc as one text response
#         text = extract_text_from_docx(file)
#         return pd.DataFrame([{"response_id": 1, "text_response": text}])

#     # Use the largest table (most rows) as the evaluation data table
#     main_table = max(tables, key=len)

#     # Standardize column names to lowercase and snake_case
#     main_table.columns = [
#         col.lower().replace(" ", "_").replace("-", "_")
#         for col in main_table.columns
#     ]

#     return main_table


# ─────────────────────────────────────────────
#  UNIFIED LOADER
# ─────────────────────────────────────────────

def load_file(file, filename: str) -> dict:
    """
    Unified entry point — detects file type from filename extension
    and routes to the correct extractor.

    Parameters
    ----------
    file     : file-like object (from Streamlit st.file_uploader)
    filename : str — the original filename, used to detect extension

    Returns
    -------
    dict with keys:
        "raw_text"   : full extracted text as one string
        "dataframe"  : pd.DataFrame if structured data found, else None
        "file_type"  : "pdf" | "docx" | "csv" | "xlsx"
        "page_count" : int (PDF only) or None
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file)
        return {
            "raw_text": text,
            "dataframe": None,       # PDFs rarely have structured table data
            "file_type": "pdf",
            "page_count": text.count("[Page ")
        }

    elif ext == ".docx":
        text = extract_text_from_docx(file)
        df = extract_structured_responses_from_docx(file)
        return {
            "raw_text": text,
            "dataframe": df,
            "file_type": "docx",
            "page_count": None
        }

    elif ext == ".csv":
        df = pd.read_csv(file)
        return {
            "raw_text": None,
            "dataframe": df,
            "file_type": "csv",
            "page_count": None
        }

    elif ext in (".xlsx", ".xls"):
        df = pd.read_excel(file)
        return {
            "raw_text": None,
            "dataframe": df,
            "file_type": "xlsx",
            "page_count": None
        }

    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: PDF, DOCX, CSV, XLSX")